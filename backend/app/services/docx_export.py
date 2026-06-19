import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.schemas.report import ReportPayload
from app.services.report_fields import build_report_rows


def render_report_docx(payload: ReportPayload, ai_summary: str | None = None) -> bytes:
    document = Document()

    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand.add_run("INFINITE")
    brand_run.bold = True
    brand_run.font.size = Pt(24)
    brand_run.font.color.rgb = RGBColor(0x6F, 0x3B, 0xD6)

    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.add_run("Managed by MEDELITE")
    tagline_run.bold = True
    tagline_run.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("FACILITY ASSESSMENT SNAPSHOT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    state = document.add_paragraph()
    state.alignment = WD_ALIGN_PARAGRAPH.CENTER
    state.add_run(payload.facility.state).bold = True

    document.add_paragraph()

    rows = build_report_rows(payload)
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[i].cells
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        value_cell.paragraphs[0].add_run(value).italic = True

    if ai_summary:
        document.add_paragraph()
        heading = document.add_paragraph()
        heading.add_run("AI-Generated Insights").bold = True
        document.add_paragraph(ai_summary)

    document.add_paragraph()
    link_paragraph = document.add_paragraph()
    link_paragraph.add_run(payload.facility.medicare_care_compare_url)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
